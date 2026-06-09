# app/crawler/parsers/form_parser.py
"""
Tải và phân tích nội dung biểu mẫu hành chính (DOCX, PDF).
Mục tiêu: extract các trường cần điền để AI có thể hướng dẫn người dùng.

Hỗ trợ:
  - .docx / .doc  → python-docx
  - .pdf          → pdfplumber
  - .xls / .xlsx  → openpyxl (biểu mẫu dạng bảng)
"""
import io
import re
import tempfile
from pathlib import Path
from typing import Any

import httpx
from loguru import logger


BASE_URL = "https://dichvucong.gov.vn"


# Tokens tiếng Việt phổ biến — biểu mẫu hành chính nào cũng có vài từ trong đây.
# Dùng để detect decode .doc OLE thành công vs rác (CJK / control chars trộn space).
_COMMON_VN_TOKENS = {
    # Pronoun / preposition / conjunction
    "của", "và", "trong", "với", "đến", "có", "không", "để", "tại",
    "này", "cho", "là", "được", "theo", "trên", "khi", "nếu", "hoặc",
    # Time / date
    "ngày", "tháng", "năm", "giờ",
    # Person / contact
    "tên", "địa", "chỉ", "người", "ông", "bà", "anh", "chị",
    # Administrative
    "thông", "tin", "thủ", "tục", "hồ", "sơ", "đăng", "ký", "giấy",
    "phép", "chứng", "minh", "nhận", "yêu", "cầu", "điều", "khoản",
    "việt", "nam", "công", "dân", "quận", "huyện", "phường", "xã",
    "tỉnh", "thành", "phố", "số", "nhà", "đường", "cộng", "hòa",
    # Form-specific
    "khai", "đơn", "mẫu", "biểu", "ban", "hành",
}


def _looks_like_vn_text(text: str, min_word_hits: int = 3) -> bool:
    """
    Heuristic: text decoded có phải tiếng Việt chuẩn (đầy đủ dấu) không?

    Check 2 điều kiện:
      1. Có ít nhất `min_word_hits` từ tiếng Việt phổ biến (token-level).
      2. Có ít nhất 5 ký tự dấu thanh VN trong vùng Latin Extended Additional
         (0x1E00-0x1EFF) — đảm bảo encoding đúng, không bị mất dấu (vd convert
         .doc legacy qua LibreOffice trên Windows mất dấu → "CỘNG" → "C?NG").

    Nếu encode đúng nhưng mất dấu (vd text ascii thuần kiểu "Cong hoa xa hoi")
    → cũng reject vì LLM không dùng được cho hướng dẫn điền form VN.
    """
    import re

    if not text or len(text) < 50:
        return False
    # ĐK 1: từ VN phổ biến
    tokens = re.findall(r"[A-Za-zÀ-ỹ]+", text.lower())
    if len(tokens) < 5:
        return False
    hits = sum(1 for t in tokens if t in _COMMON_VN_TOKENS)
    if hits < min_word_hits:
        return False
    # ĐK 2: dấu thanh VN (sample đầu 3000 char đủ)
    diacritic_count = sum(
        1 for ch in text[:3000] if 0x1E00 <= ord(ch) <= 0x1EFF
    )
    return diacritic_count >= 5


class FormField:
    """Một trường trong biểu mẫu."""
    def __init__(self, label: str, hint: str = "", required: bool = True):
        self.label = label.strip()
        self.hint = hint.strip()
        self.required = required

    def __repr__(self):
        return f"FormField({self.label!r})"


class FormParser:
    """
    Download biểu mẫu từ URL → extract danh sách trường cần điền.
    """

    async def parse_form(self, form_url: str, form_name: str = "") -> dict[str, Any] | None:
        """
        Download (GET) và parse biểu mẫu — convenience wrapper.

        Lưu ý: DVCQG đòi POST cho /preview-attachment nên đường GET sẽ trả
        "Request Rejected". Crawler thật dùng `download_attachment()` + truyền
        bytes vào `parse_bytes()` trực tiếp; `parse_form()` chỉ tiện cho test
        host khác hoặc URL legacy.
        """
        try:
            content, filename, content_type = await self._download(form_url)
            if not content:
                return None
            result = self.parse_bytes(
                content=content,
                form_name=form_name,
                filename_hint=filename,
                content_type=content_type,
                source_url=form_url,
            )
            if result is None:
                return None
            result["form_url"] = form_url
            return result
        except Exception as e:
            logger.error(f"FormParser | failed | url={form_url} | error={e}")
            return None

    def parse_bytes(
        self,
        content: bytes,
        form_name: str = "",
        filename_hint: str = "",
        content_type: str = "",
        source_url: str = "",
    ) -> dict[str, Any] | None:
        """
        Parse bytes đã download sẵn → fields + raw_text.

        Caller (vd crawler) tự handle HTTP — DVCQG cần POST `/preview-attachment`
        chứ không GET được. Trả về None nếu không extract được gì có nghĩa.

        Status để caller persist:
          - 'ok'          → có raw_text hoặc fields
          - 'unsupported' → ext không nhận diện được, decode fail
          - 'failed'      → exception khi parse
        """
        try:
            ext = self._detect_extension(filename_hint or form_name, content_type, source_url)
            logger.info(
                f"FormParser | parsing | name={form_name} | ext={ext} | size={len(content)}B"
            )

            if ext == "docx":
                fields, raw_text = self._parse_docx(content)
            elif ext == "doc":
                fields, raw_text = self._parse_doc_binary(content)
            elif ext == "pdf":
                fields, raw_text = self._parse_pdf(content)
            elif ext in ("xlsx", "xls"):
                fields, raw_text = self._parse_excel(content)
            else:
                # Fallback: parse như plain text
                raw_text = content.decode("utf-8", errors="ignore")
                fields = self._extract_fields_from_text(raw_text)
                if not raw_text.strip():
                    return {
                        "form_name": form_name,
                        "fields": [],
                        "raw_text": "",
                        "ext": ext,
                        "status": "unsupported",
                    }

            if not fields and not raw_text:
                logger.warning(f"FormParser | no content extracted | name={form_name}")
                return {
                    "form_name": form_name,
                    "fields": [],
                    "raw_text": "",
                    "ext": ext,
                    "status": "unsupported",
                }

            return {
                "form_name": form_name,
                "fields": [
                    {"label": f.label, "hint": f.hint, "required": f.required}
                    for f in fields
                ],
                "raw_text": raw_text[:5000],
                "ext": ext,
                "status": "ok",
            }
        except Exception as e:
            logger.error(f"FormParser | parse_bytes failed | name={form_name} | {e}")
            return {
                "form_name": form_name,
                "fields": [],
                "raw_text": "",
                "ext": "unknown",
                "status": "failed",
            }

    # ── Download ──────────────────────────────────────────────────────────────

    async def _download(self, url: str) -> tuple[bytes | None, str, str]:
        """Tải file, trả về (content, filename, content_type)."""
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Referer": BASE_URL,
        }
        try:
            async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
                resp = await client.get(url, headers=headers)
                resp.raise_for_status()

                content_type = resp.headers.get("content-type", "")
                # Lấy filename từ Content-Disposition hoặc URL
                cd = resp.headers.get("content-disposition", "")
                filename = ""
                if cd:
                    # RFC 5987: filename*=UTF-8''encoded_name
                    import urllib.parse
                    m = re.search(r"filename\*=(?:UTF-8'')?([^\s;]+)", cd, re.IGNORECASE)
                    if m:
                        filename = urllib.parse.unquote(m.group(1)).strip().strip('"\'')
                    # RFC 2183: filename="name"
                    if not filename:
                        m = re.search(r'filename=["\']?([^"\';\r\n]+)["\']?', cd, re.IGNORECASE)
                        if m:
                            filename = m.group(1).strip().strip('"\'')
                if not filename:
                    filename = url.split("/")[-1].split("?")[0]

                return resp.content, filename, content_type
        except Exception as e:
            logger.warning(f"FormParser | download failed | url={url} | {e}")
            return None, "", ""

    def _detect_extension(self, filename: str, content_type: str, url: str) -> str:
        """Xác định loại file từ filename, content-type hoặc URL."""
        for source in [filename, url]:
            ext = Path(source.split("?")[0]).suffix.lower().lstrip(".")
            if ext in ("docx", "doc", "pdf", "xlsx", "xls"):
                return ext

        ct = content_type.lower()
        if "pdf" in ct:
            return "pdf"
        if "word" in ct or "docx" in ct or "officedocument" in ct:
            return "docx"
        if "excel" in ct or "spreadsheet" in ct:
            return "xlsx"

        return "unknown"

    # ── DOCX Parser ───────────────────────────────────────────────────────────

    def _parse_docx(self, content: bytes) -> tuple[list[FormField], str]:
        """Extract các trường từ file DOCX."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
        except Exception as e:
            logger.warning(f"FormParser | docx open error | {e}")
            return [], ""

        raw_lines = []
        fields = []

        # 1. Đọc tất cả paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                raw_lines.append(text)

        # 2. Đọc tất cả tables — thường form có dạng bảng label | value
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells:
                    continue

                raw_lines.append(" | ".join(cells))

                # Phát hiện trường: ô đầu là label, ô sau là blank (chỗ điền)
                if len(cells) >= 1:
                    label = cells[0]
                    # Lọc label hợp lệ (không phải header, không quá dài)
                    if self._is_valid_field_label(label):
                        hint = self._generate_hint(label)
                        fields.append(FormField(label=label, hint=hint))

        # 3. Extract thêm từ text nếu ít fields
        if len(fields) < 3:
            text_fields = self._extract_fields_from_text("\n".join(raw_lines))
            # Merge, tránh duplicate
            existing = {f.label for f in fields}
            for f in text_fields:
                if f.label not in existing:
                    fields.append(f)

        raw_text = "\n".join(raw_lines)
        return fields, raw_text

    # ── DOC Binary Parser (OLE Word 97-2003) ─────────────────────────────────

    def _parse_doc_binary(self, content: bytes) -> tuple[list[FormField], str]:
        """
        Parse .doc OLE Compound Binary File (Word 97-2003).

        Strategy:
          1. Thử python-docx trước (một số .doc thực ra là .docx rename).
          2. Thử LibreOffice/antiword qua subprocess nếu binary có trên PATH —
             đây là cách RELIABLE duy nhất parse .doc.
          3. Bỏ. Return empty → caller set status='unsupported'.

        Tại sao KHÔNG decode UTF-16 LE thẳng: .doc OLE Compound File lưu text
        rải rác qua nhiều stream (WordDocument, 1Table, Pieces), encode phụ
        thuộc FIB header. Decode raw UTF-16 LE chỉ tóm được 1 phần, output
        trộn lẫn text VN + rác (CJK/control) — không dùng được cho LLM.

        Để support .doc đầy đủ trên prod:
          - Debian/Ubuntu: apt-get install libreoffice-core libreoffice-writer
            (hoặc nhẹ hơn: apt-get install antiword)
          - Windows dev: cài LibreOffice + thêm vào PATH
        """
        # 1. Thử docx (file rename .doc nhưng thực là .docx)
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            raw_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        raw_lines.append(" | ".join(cells))
            if raw_lines:
                text = "\n".join(raw_lines)
                fields = self._extract_fields_from_text(text)
                return fields, text
        except Exception:
            pass

        # 2. LibreOffice headless / antiword
        text = self._convert_doc_via_external(content)
        if text and _looks_like_vn_text(text, min_word_hits=3):
            fields = self._extract_fields_from_text(text)
            logger.info(
                f"FormParser | doc via external | text_len={len(text)} "
                f"| fields={len(fields)}"
            )
            return fields, text[:5000]

        logger.warning(
            "FormParser | .doc parse skipped — install libreoffice-headless "
            "hoặc antiword để support file Word 97-2003"
        )
        return [], ""

    def _convert_doc_via_external(self, content: bytes) -> str:
        """Convert .doc → text qua LibreOffice hoặc antiword (best-effort)."""
        import shutil
        import subprocess
        import tempfile

        # antiword — nhanh, nhẹ, output text
        antiword = shutil.which("antiword")
        if antiword:
            try:
                with tempfile.NamedTemporaryFile(
                    suffix=".doc", delete=False
                ) as tf:
                    tf.write(content)
                    tf_path = tf.name
                try:
                    result = subprocess.run(
                        [antiword, tf_path],
                        capture_output=True, timeout=15,
                    )
                    if result.returncode == 0:
                        return result.stdout.decode("utf-8", errors="ignore")
                finally:
                    Path(tf_path).unlink(missing_ok=True)
            except Exception as e:
                logger.warning(f"FormParser | antiword failed | {e}")

        # LibreOffice headless — convert .doc → .docx → đọc bằng python-docx
        # (chính xác hơn convert-to-txt vì txt filter bỏ table/encoding tệ).
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            try:
                with tempfile.TemporaryDirectory() as td:
                    in_path = Path(td) / "in.doc"
                    in_path.write_bytes(content)
                    subprocess.run(
                        [
                            soffice, "--headless",
                            "--convert-to", "docx",
                            "--outdir", td, str(in_path),
                        ],
                        capture_output=True, timeout=45,
                    )
                    out_path = Path(td) / "in.docx"
                    if out_path.exists():
                        docx_bytes = out_path.read_bytes()
                        from docx import Document
                        doc = Document(io.BytesIO(docx_bytes))
                        lines = []
                        for p in doc.paragraphs:
                            t = p.text.strip()
                            if t:
                                lines.append(t)
                        for table in doc.tables:
                            for row in table.rows:
                                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                                if cells:
                                    lines.append(" | ".join(cells))
                        return "\n".join(lines)
            except Exception as e:
                logger.warning(f"FormParser | libreoffice doc→docx failed | {e}")

        return ""

    # ── PDF Parser ────────────────────────────────────────────────────────────

    def _parse_pdf(self, content: bytes) -> tuple[list[FormField], str]:
        """Extract các trường từ file PDF."""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("FormParser | pdfplumber not installed | pip install pdfplumber")
            return [], ""

        raw_lines = []
        fields = []

        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    # Extract text
                    text = page.extract_text() or ""
                    if text:
                        raw_lines.extend(text.split("\n"))

                    # Extract bảng
                    for table in page.extract_tables() or []:
                        for row in table:
                            cells = [str(c).strip() for c in row if c and str(c).strip()]
                            if cells:
                                raw_lines.append(" | ".join(cells))
                                label = cells[0]
                                if self._is_valid_field_label(label):
                                    hint = self._generate_hint(label)
                                    fields.append(FormField(label=label, hint=hint))

                    # Extract fillable form fields (nếu PDF có form)
                    if hasattr(page, 'annots') and page.annots:
                        for annot in page.annots:
                            if annot.get("subtype") == "Widget":
                                label = annot.get("T", "") or annot.get("TU", "")
                                if label and self._is_valid_field_label(label):
                                    fields.append(FormField(
                                        label=label,
                                        hint=self._generate_hint(label)
                                    ))

        except Exception as e:
            logger.warning(f"FormParser | pdf parse error | {e}")

        # Fallback extract từ text nếu ít fields
        if len(fields) < 3:
            text_fields = self._extract_fields_from_text("\n".join(raw_lines))
            existing = {f.label for f in fields}
            for f in text_fields:
                if f.label not in existing:
                    fields.append(f)

        return fields, "\n".join(raw_lines[:200])  # giới hạn raw text

    # ── Excel Parser ──────────────────────────────────────────────────────────

    def _parse_excel(self, content: bytes) -> tuple[list[FormField], str]:
        """Extract các trường từ file Excel."""
        try:
            import openpyxl
        except ImportError:
            logger.warning("FormParser | openpyxl not installed | pip install openpyxl")
            return [], ""

        raw_lines = []
        fields = []

        try:
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if not cells:
                        continue
                    raw_lines.append(" | ".join(cells))
                    label = cells[0]
                    if self._is_valid_field_label(label):
                        hint = self._generate_hint(label)
                        fields.append(FormField(label=label, hint=hint))
        except Exception as e:
            logger.warning(f"FormParser | excel parse error | {e}")

        return fields, "\n".join(raw_lines)

    # ── Field extraction từ plain text ───────────────────────────────────────

    def _extract_fields_from_text(self, text: str) -> list[FormField]:
        """
        Tìm các trường dạng:
          - "Họ và tên: .................."
          - "1. Số CMND/CCCD:"
          - "[Họ tên]"
        """
        fields = []
        seen = set()

        patterns = [
            r'^(\d+[\.\)]\s*)?([A-ZÀ-Ỹa-zà-ỹ][^:\n]{3,60})\s*:\s*[\.\_\s]{3,}',  # "Họ tên: ....."
            r'^(\d+[\.\)]\s*)?([A-ZÀ-Ỹa-zà-ỹ][^:\n]{3,60})\s*:\s*$',               # "Họ tên:"
            r'\[([A-ZÀ-Ỹa-zà-ỹ][^\]]{3,60})\]',                                      # "[Họ tên]"
        ]

        for line in text.split("\n"):
            line = line.strip()
            if not line or len(line) > 200:
                continue

            for pattern in patterns:
                m = re.search(pattern, line)
                if m:
                    label = m.group(2) if len(m.groups()) >= 2 else m.group(1)
                    label = label.strip().rstrip(":")
                    if label and label not in seen and self._is_valid_field_label(label):
                        seen.add(label)
                        fields.append(FormField(
                            label=label,
                            hint=self._generate_hint(label)
                        ))
                    break

        return fields

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _is_valid_field_label(self, label: str) -> bool:
        """Kiểm tra label có phải là trường hợp lệ không."""
        if not label or len(label) < 3 or len(label) > 150:
            return False
        # Loại bỏ header, số thứ tự, text quá ngắn hoặc toàn số
        skip_patterns = [
            r'^\d+$',                          # toàn số
            r'^(STT|TT|No\.|Số thứ tự)$',     # số thứ tự
            r'^(Ghi chú|Note|Chú ý)$',         # ghi chú đơn thuần
        ]
        for p in skip_patterns:
            if re.match(p, label, re.IGNORECASE):
                return False
        return True

    def _generate_hint(self, label: str) -> str:
        """
        Tạo gợi ý điền dựa trên tên trường.
        Dùng rule-based cho các trường phổ biến.
        """
        label_lower = label.lower()

        hints = {
            ("họ và tên", "họ tên", "tên đầy đủ"):
                "Ghi đầy đủ họ, chữ đệm và tên theo CMND/CCCD (chữ in hoa)",
            ("ngày sinh", "năm sinh", "ngày tháng năm sinh"):
                "Ghi theo định dạng DD/MM/YYYY (ví dụ: 15/08/1990)",
            ("giới tính",):
                "Đánh dấu X vào ô Nam hoặc Nữ",
            ("số cmnd", "số cccd", "số căn cước", "chứng minh nhân dân"):
                "Ghi số CCCD 12 chữ số hoặc CMND 9 chữ số",
            ("nơi sinh", "tỉnh/thành phố sinh"):
                "Ghi tên tỉnh/thành phố nơi bạn sinh ra",
            ("hộ khẩu thường trú", "thường trú", "địa chỉ thường trú"):
                "Ghi đầy đủ số nhà, đường/thôn, xã/phường, huyện/quận, tỉnh/thành phố",
            ("tạm trú", "địa chỉ tạm trú"):
                "Ghi địa chỉ nơi đang tạm trú hiện tại (nếu có)",
            ("dân tộc",):
                "Ghi tên dân tộc (ví dụ: Kinh, Tày, Mường...)",
            ("quốc tịch",):
                "Ghi quốc tịch (ví dụ: Việt Nam)",
            ("tôn giáo",):
                "Ghi tên tôn giáo hoặc để trống nếu không theo tôn giáo nào",
            ("trình độ học vấn", "học vấn"):
                "Ghi cấp học cao nhất đã hoàn thành",
            ("nghề nghiệp",):
                "Ghi nghề nghiệp hiện tại",
            ("điện thoại", "số điện thoại", "điện thoại liên hệ"):
                "Ghi số điện thoại liên hệ (10 chữ số)",
            ("email",):
                "Ghi địa chỉ email (nếu có)",
            ("ngày cấp",):
                "Ghi ngày cấp CMND/CCCD theo định dạng DD/MM/YYYY",
            ("nơi cấp",):
                "Ghi cơ quan cấp CMND/CCCD (ví dụ: Công an TP. Hà Nội)",
            ("họ tên cha", "tên cha"):
                "Ghi đầy đủ họ tên người cha",
            ("họ tên mẹ", "tên mẹ"):
                "Ghi đầy đủ họ tên người mẹ",
            ("họ tên vợ", "họ tên chồng", "vợ/chồng"):
                "Ghi đầy đủ họ tên vợ hoặc chồng",
            ("ngày đăng ký", "ngày làm thủ tục"):
                "Hệ thống tự điền hoặc ghi ngày nộp hồ sơ",
            ("chữ ký", "ký tên"):
                "Ký tên trực tiếp vào ô này",
        }

        for keywords, hint in hints.items():
            if any(kw in label_lower for kw in keywords):
                return hint

        # Fallback chung
        return f"Điền thông tin về: {label}"


def format_form_chunk(form_data: dict, procedure_name: str) -> str:
    """
    Tạo nội dung chunk từ form đã parse — dùng để embed vào ChromaDB.
    AI sẽ dùng chunk này để hướng dẫn người dùng điền biểu mẫu.
    """
    lines = [
        f"Thủ tục: {procedure_name}",
        f"Biểu mẫu: {form_data['form_name']}",
        f"Link tải: {form_data['form_url']}",
        "",
        "Hướng dẫn điền các trường:",
    ]

    fields = form_data.get("fields", [])
    if fields:
        for i, field in enumerate(fields, 1):
            line = f"{i}. {field['label']}"
            if field.get("hint"):
                line += f": {field['hint']}"
            lines.append(line)
    else:
        # Không extract được fields → dùng raw text
        raw = form_data.get("raw_text", "").strip()
        if raw:
            lines.append("Nội dung biểu mẫu:")
            lines.append(raw[:2000])

    return "\n".join(lines)
